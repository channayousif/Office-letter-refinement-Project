"""Utility functions for handling .docx files."""
from typing import Tuple, List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

class DocumentElement:
    """Base class for document elements."""
    def __init__(self, element_type: str):
        self.element_type = element_type

class TextElement(DocumentElement):
    """Class representing a text element in a document."""
    def __init__(self, text: str):
        super().__init__("text")
        self.text = text

class TableElement(DocumentElement):
    """Class representing a table element in a document."""
    def __init__(self, rows: List[List[str]], columns: List[str]):
        super().__init__("table")
        self.rows = rows
        self.columns = columns

class ImageElement(DocumentElement):
    """Class representing an image element in a document."""
    def __init__(self, image_path: str, caption: Optional[str] = None):
        super().__init__("image")
        self.image_path = image_path
        self.caption = caption

def extract_document_elements(file_path: str) -> List[DocumentElement]:
    """
    Extract elements from a .docx file.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        List[DocumentElement]: List of document elements
    """
    doc = Document(file_path)
    elements = []
    
    # Process paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            elements.append(TextElement(para.text))
    
    # Process tables
    for table in doc.tables:
        # Check if table has content
        if len(table.rows) == 0:
            continue
            
        # Get information about merged cells
        # We need to track which cells are part of merged regions
        # This is complex in python-docx, so we'll use a simpler approach
        # by tracking cells that appear to have the same content
        
        # First, extract raw cell data
        raw_rows = []
        for row in table.rows:
            raw_cells = []
            for cell in row.cells:
                # Clean cell text
                cell_text = cell.text.strip() if cell.text else ""
                raw_cells.append(cell_text)
            raw_rows.append(raw_cells)
        
        # Now process the data to handle merged cells
        processed_rows = []
        for row_idx, raw_row in enumerate(raw_rows):
            processed_row = []
            skip_count = 0
            
            for col_idx, cell_text in enumerate(raw_row):
                # Skip cells that we've identified as part of a merge
                if skip_count > 0:
                    skip_count -= 1
                    continue
                    
                # Check if this cell is repeated (potential horizontal merge)
                repeated = False
                repeat_count = 0
                
                if col_idx < len(raw_row) - 1:
                    next_idx = col_idx + 1
                    while next_idx < len(raw_row) and raw_row[next_idx] == cell_text and cell_text:
                        repeat_count += 1
                        next_idx += 1
                
                # Only consider it a merged cell if repeated more than once and not empty
                if repeat_count > 0 and cell_text:
                    repeated = True
                    skip_count = repeat_count
                
                # Add the cell text
                processed_row.append(cell_text)
                
            processed_rows.append(processed_row)
        
        # Extract column headers from first row
        columns = []
        if len(processed_rows) > 0:
            first_row = processed_rows[0]
            for i, cell_text in enumerate(first_row):
                # Handle empty column names
                header = cell_text.strip() if cell_text else f"Column_{i+1}"
                columns.append(header)
        
        # Process column names to make them unique
        processed_columns = []
        column_counts = {}
        
        for col in columns:
            if col in column_counts:
                column_counts[col] += 1
                processed_columns.append(f"{col}_{column_counts[col]}")
            else:
                column_counts[col] = 0
                processed_columns.append(col)
        
        # Ensure all rows have the same number of columns
        max_cols = len(processed_columns)
        for i in range(len(processed_rows)):
            while len(processed_rows[i]) < max_cols:
                processed_rows[i].append("")
        
        # Create TableElement with processed column names
        table_element = TableElement(processed_rows, processed_columns)
        
        # Add the table to the elements list
        elements.append(table_element)
    
    return elements

def create_docx_from_elements(elements: List[DocumentElement], output_path: str) -> None:
    """
    Create a new .docx file from document elements.
    
    Args:
        elements: List of document elements
        output_path: Path where to save the .docx file
    """
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add each element to the document
    for element in elements:
        if isinstance(element, TextElement):
            if element.text.strip():
                # Create a new paragraph with the text
                # Split by newlines to preserve paragraph structure
                for para_text in element.text.split('\n\n'):
                    if para_text.strip():
                        para = doc.add_paragraph()
                        # Add text without any markdown formatting
                        para.add_run(para_text.strip())
                
        elif isinstance(element, TableElement):
            # Create table with proper dimensions
            if len(element.rows) > 0 and len(element.columns) > 0:
                # Add a small space before table
                doc.add_paragraph()
                
                # Create the table with the correct number of rows and columns
                num_rows = len(element.rows)
                num_cols = len(element.columns)
                
                # Ensure we have at least one row and column
                if num_rows == 0 or num_cols == 0:
                    continue
                    
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'  # Add borders
                
                # Track cells to merge
                merged_cells = []
                
                # Add header row in bold
                for j, col_name in enumerate(element.columns):
                    if j < num_cols:  # Ensure we don't go out of bounds
                        cell = table.cell(0, j)
                        # Clean column name
                        clean_col_name = col_name.strip()
                        # Remove any numeric suffixes added for uniqueness
                        import re
                        clean_col_name = re.sub(r'_\d+$', '', clean_col_name)
                        run = cell.paragraphs[0].add_run(clean_col_name)
                        run.bold = True  # Make header bold
                
                # Add data rows and detect potential merged cells
                for i in range(1, num_rows):  # Skip header row, start from 1
                    row = element.rows[i]
                    
                    # First pass: add content to cells
                    for j in range(num_cols):
                        if j < len(row):
                            cell = table.cell(i, j)
                            # Clean any remaining markdown from cell text
                            cell_text = row[j] if j < len(row) else ""
                            clean_cell_text = cell_text.strip() if cell_text else ""
                            import re
                            clean_cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_cell_text)
                            clean_cell_text = re.sub(r'\*(.*?)\*', r'\1', clean_cell_text)
                            cell.text = clean_cell_text
                
                # Try to identify and merge horizontal cell spans
                for i in range(num_rows):
                    j = 0
                    while j < num_cols - 1:
                        # Check if this cell and next cell have identical content
                        current_cell = table.cell(i, j)
                        start_j = j
                        
                        # Count how many consecutive cells have the same content
                        span_length = 1
                        while j + span_length < num_cols:
                            next_cell = table.cell(i, j + span_length)
                            if current_cell.text == next_cell.text and current_cell.text.strip():
                                span_length += 1
                            else:
                                break
                        
                        # If we found a span, merge the cells
                        if span_length > 1:
                            try:
                                # Create a merged cell
                                table.cell(i, start_j).merge(table.cell(i, start_j + span_length - 1))
                                j += span_length
                            except Exception:
                                # If merge fails, just move to the next cell
                                j += 1
                        else:
                            j += 1
                
                # Add spacing after table
                doc.add_paragraph()
                
        elif isinstance(element, ImageElement):
            # Handle images if needed
            if hasattr(element, 'image_path') and element.image_path:
                try:
                    doc.add_picture(element.image_path, width=Inches(6))
                    if element.caption:
                        caption = doc.add_paragraph(element.caption)
                        caption.style = 'Caption'
                except Exception:
                    # If image can't be added, add a placeholder
                    doc.add_paragraph("[Image placeholder]")
    
    # Save the document
    doc.save(output_path)

def read_docx(file_path: str) -> str:
    """
    Read content from a .docx file.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        str: Extracted text content
    """
    doc = Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    return "\n\n".join(full_text)

def create_docx(content: str, output_path: str) -> None:
    """
    Create a new .docx file with the given content.
    
    Args:
        content: Text content to write
        output_path: Path where to save the .docx file
    """
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Split content into paragraphs
    paragraphs = content.split("\n\n")
    
    # Add each paragraph to the document
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
    
    # Save the document
    doc.save(output_path)

def compare_versions(original: str, refined: str) -> Tuple[str, List[Dict]]:
    """
    Compare original and refined versions of the text.
    
    Args:
        original: Original text content
        refined: Refined text content
        
    Returns:
        Tuple[str, list]: Summary of changes and list of specific modifications
    """
    # Split into paragraphs
    original_paras = original.split("\n\n")
    refined_paras = refined.split("\n\n")
    
    changes = []
    summary_points = []
    
    # Compare paragraphs
    for i, (orig, ref) in enumerate(zip(original_paras, refined_paras)):
        if orig.strip() != ref.strip():
            changes.append({
                "paragraph": i + 1,
                "original": orig.strip(),
                "refined": ref.strip()
            })
            
            # Generate a summary point for this change
            if len(orig) != len(ref):
                if len(ref) > len(orig):
                    summary_points.append(f"Paragraph {i + 1}: Expanded for clarity and detail")
                else:
                    summary_points.append(f"Paragraph {i + 1}: Condensed for conciseness")
            else:
                summary_points.append(f"Paragraph {i + 1}: Refined for improved clarity and professionalism")
    
    # Handle case where no changes were made
    if not summary_points:
        summary = "No significant changes were required. The document already meets professional standards."
    else:
        summary = "The following improvements were made:\n- " + "\n- ".join(summary_points)
    
    return summary, changes 